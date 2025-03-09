# import re
# import json
# import pdfplumber

# text = ""
# with pdfplumber.open("31-2024-qh15_1.pdf") as pdf:
#     for page in pdf.pages:
#         text += page.extract_text() + "\n"

# # Khởi tạo cấu trúc cơ bản của văn bản luật
# law = {
#     "id": "law_001",
#     "title": "LUẬT ĐẤT ĐAI",
#     "content": text,   # Nội dung đầy đủ (bạn có thể bỏ nếu không cần lưu)
#     "date": "01-03-2024",
#     "source": "Quốc hội",
#     "category": "Đất đai",
#     "chapters": []
# }

# # Định nghĩa biểu thức chính quy để nhận diện Chương, Mục, Điều và Điều khoản
# chapter_pattern = re.compile(r"^Chương\s+([IVXLCDM]+)(.*)", re.IGNORECASE)
# section_pattern = re.compile(r"^Mục\s+(\d+)\s*(.*)", re.IGNORECASE)
# article_pattern = re.compile(r"^Điều\s+(\d+)[\.\:]\s*(.*)", re.IGNORECASE)
# # Pattern cho điều khoản, giả sử bắt đầu bằng số và dấu chấm (ví dụ: "1. Nội dung điều khoản")
# clause_pattern = re.compile(r"^(\d+)\.\s*(.*)")

# lines = text.splitlines()
# i = 0
# current_chapter = None
# current_section = None
# current_article = None
# current_clause = None

# while i < len(lines):
#     line = lines[i].strip()
#     if not line:
#         i += 1
#         continue

#     # Kiểm tra tiêu đề Chương
#     chapter_match = chapter_pattern.match(line)
#     if chapter_match:
#         roman_num = chapter_match.group(1)
#         chapter_title = chapter_match.group(2).strip()
#         current_chapter = {
#             "id": f"chapter_{roman_num}",
#             "name": f"Chương {roman_num} {chapter_title}",
#             "number": roman_num,
#             "content": "",
#             "articles": []
#         }
#         law["chapters"].append(current_chapter)
#         current_section = None
#         current_article = None
#         current_clause = None
#         i += 1
#         continue

#     # Kiểm tra tiêu đề Mục
#     section_match = section_pattern.match(line)
#     if section_match and current_chapter:
#         section_number = section_match.group(1)
#         section_title = section_match.group(2).strip()
#         full_section_header = f"Mục {section_number} {section_title}".strip()
#         j = i + 1
#         while j < len(lines):
#             next_line = lines[j].strip()
#             if not next_line:
#                 j += 1
#                 continue
#             if chapter_pattern.match(next_line) or article_pattern.match(next_line) or section_pattern.match(next_line):
#                 break
#             full_section_header += " " + next_line
#             j += 1
#         current_section = full_section_header
#         i = j
#         continue

#     # Kiểm tra tiêu đề Điều
#     article_match = article_pattern.match(line)
#     if article_match and current_chapter:
#         article_number = article_match.group(1)
#         article_rest = article_match.group(2).strip()
#         current_article = {
#             "id": f"article_{article_number}",
#             "parent": current_section if current_section is not None else None,
#             "name": article_rest,  # Ở đây có thể tách riêng nếu cần
#             "number": int(article_number),
#             "content": article_rest,
#             "clauses": []  # Danh sách chứa các điều khoản của điều này
#         }
#         current_chapter["articles"].append(current_article)
#         current_clause = None  # Reset điều khoản cho điều mới
#         i += 1
#         continue

#     # Kiểm tra nếu dòng là tiêu đề của điều khoản (clause)
#     clause_match = clause_pattern.match(line)
#     if clause_match and current_article:
#         clause_number = clause_match.group(1)
#         clause_text = clause_match.group(2).strip()
#         current_clause = {
#             "id": f"clause_{current_article['number']}_{clause_number}",
#             "number": int(clause_number),
#             "content": clause_text
#         }
#         current_article["clauses"].append(current_clause)
#         i += 1
#         continue

#     # Nếu dòng không khớp với tiêu đề nào, gán vào nội dung của điều khoản nếu đang có, nếu không thì gán vào nội dung của điều
#     if current_clause:
#         current_clause["content"] += " " + line
#     elif current_article:
#         current_article["content"] += " " + line
#     elif current_chapter:
#         current_chapter["content"] += " " + line

#     i += 1

# # Lưu kết quả ra file JSON
# with open("law_structure.json", "w", encoding="utf-8") as f:
#     json.dump({"law": law}, f, ensure_ascii=False, indent=4)
















# ###################################################################
# import re
# import json
# from docx import Document

# # Đọc dữ liệu từ file DOCX
# doc = Document("31-2024-qh15_1.docx")
# # Lấy danh sách các đoạn văn mà không nối bằng "\n"
# lines = [para.text for para in doc.paragraphs]

# # Khởi tạo cấu trúc cơ bản của văn bản luật (bỏ thuộc tính "content")
# law = {
#     "id": "law_001",
#     "title": "LUẬT ĐẤT ĐAI",
#     "date": "01-03-2024",
#     "source": "Quốc hội",
#     "category": "Đất đai",
#     "chapters": []
# }

# # Định nghĩa các biểu thức chính quy để nhận diện Chương, Mục, Điều và Điều khoản
# chapter_pattern = re.compile(r"^Chương\s+([IVXLCDM]+)(.*)", re.IGNORECASE)
# section_pattern = re.compile(r"^Mục\s+(\d+)\s*(.*)", re.IGNORECASE)
# article_pattern = re.compile(r"^Điều\s+(\d+)[\.\:]\s*(.*)", re.IGNORECASE)
# clause_pattern = re.compile(r"^(\d+)\.\s*(.*)")

# i = 0
# current_chapter = None
# current_section = None
# current_article = None
# current_clause = None

# while i < len(lines):
#     line = lines[i].strip()
#     if not line:
#         i += 1
#         continue

#     # Kiểm tra tiêu đề Chương
#     chapter_match = chapter_pattern.match(line)
#     if chapter_match:
#         roman_num = chapter_match.group(1)
#         chapter_title = chapter_match.group(2).strip()
#         current_chapter = {
#             "id": f"chapter_{roman_num}",
#             "name": f"Chương {roman_num} {chapter_title}",
#             "number": roman_num,
#             "content": "",
#             "articles": []
#         }
#         law["chapters"].append(current_chapter)
#         current_section = None
#         current_article = None
#         current_clause = None
#         i += 1
#         continue

#     # Kiểm tra tiêu đề Mục
#     section_match = section_pattern.match(line)
#     if section_match and current_chapter:
#         section_number = section_match.group(1)
#         section_title = section_match.group(2).strip()
#         full_section_header = f"Mục {section_number} {section_title}".strip()
#         j = i + 1
#         while j < len(lines):
#             next_line = lines[j].strip()
#             if not next_line:
#                 j += 1
#                 continue
#             if chapter_pattern.match(next_line) or article_pattern.match(next_line) or section_pattern.match(next_line):
#                 break
#             full_section_header += " " + next_line
#             j += 1
#         current_section = full_section_header
#         i = j
#         continue

#     # Kiểm tra tiêu đề Điều
#     article_match = article_pattern.match(line)
#     if article_match and current_chapter:
#         article_number = article_match.group(1)
#         article_rest = article_match.group(2).strip()
#         current_article = {
#             "id": f"article_{article_number}",
#             "parent": current_section if current_section is not None else None,
#             "name": article_rest,
#             "number": int(article_number),
#             "content": article_rest,
#             "clauses": []
#         }
#         current_chapter["articles"].append(current_article)
#         current_clause = None  # Reset cho điều mới
#         i += 1
#         continue

#     # Kiểm tra tiêu đề của điều khoản (clause)
#     clause_match = clause_pattern.match(line)
#     if clause_match and current_article:
#         clause_number = clause_match.group(1)
#         clause_text = clause_match.group(2).strip()
#         current_clause = {
#             "id": f"clause_{current_article['number']}_{clause_number}",
#             "number": int(clause_number),
#             "content": clause_text
#         }
#         current_article["clauses"].append(current_clause)
#         i += 1
#         continue

#     # Nếu dòng không khớp với tiêu đề nào, gán vào nội dung của điều khoản nếu đang có,
#     # nếu không thì gán vào nội dung của điều
#     if current_clause:
#         current_clause["content"] += " " + line
#     elif current_article:
#         current_article["content"] += " " + line
#     elif current_chapter:
#         current_chapter["content"] += " " + line

#     i += 1

# # Lưu kết quả ra file JSON
# with open("law_structure.json", "w", encoding="utf-8") as f:
#     json.dump({"law": law}, f, ensure_ascii=False, indent=4)

# print("Đã chuyển file DOC sang DOCX, đọc file DOCX và cập nhật cấu trúc văn bản luật thành công.")




############################################




# import re
# import json
# from docx import Document

# # Đọc dữ liệu từ file DOCX
# doc = Document("31-2024-qh15_1.docx")
# lines = [para.text for para in doc.paragraphs]

# # Khởi tạo cấu trúc cơ bản của văn bản luật
# law = {
#     "id": "law_001",
#     "title": "LUẬT ĐẤT ĐAI",
#     "date": "01-03-2024",
#     "source": "Quốc hội",
#     "category": "Đất đai",
#     "chapters": []
# }

# # Định nghĩa các biểu thức chính quy
# chapter_pattern = re.compile(r"^Chương\s+([IVXLCDM]+)(.*)", re.IGNORECASE)
# section_pattern = re.compile(r"^Mục\s+(\d+)\s*(.*)", re.IGNORECASE)
# article_pattern = re.compile(r"^Điều\s+(\d+)[\.\:]\s*(.*)", re.IGNORECASE)
# clause_pattern = re.compile(r"^(\d+)\.\s*(.*)")
# point_pattern = re.compile(r"^([a-z])[)]\s*(.*)")

# # Hàm chuẩn hóa tiêu đề với xử lý số La Mã
# def capitalize_title(text):
#     words = text.split()
#     for i, word in enumerate(words):
#         # Nếu từ là số La Mã (chỉ chứa I, V, X, L, C, D, M), chuyển thành in hoa
#         if re.match(r'^[IVXLCDM]+$', word, re.IGNORECASE):
#             words[i] = word.upper()
#         else:
#             words[i] = word.capitalize()
#     return " ".join(words)

# i = 0
# current_chapter = None
# current_section = None
# current_article = None
# current_clause = None
# current_point = None

# while i < len(lines):
#     line = lines[i].strip()
#     if not line:
#         i += 1
#         continue

#     # Kiểm tra tiêu đề Chương
#     chapter_match = chapter_pattern.match(line)
#     if chapter_match:
#         roman_num = chapter_match.group(1).upper()  # Chuẩn hóa số La Mã thành in hoa
#         chapter_title = chapter_match.group(2).strip()
#         full_chapter_title = f"Chương {roman_num} {chapter_title}".strip()
        
#         # Kiểm tra nếu tiêu đề chương kéo dài qua nhiều dòng
#         j = i + 1
#         while j < len(lines):
#             next_line = lines[j].strip()
#             if not next_line or chapter_pattern.match(next_line) or section_pattern.match(next_line) or article_pattern.match(next_line):
#                 break
#             full_chapter_title += " " + next_line
#             j += 1
        
#         current_chapter = {
#             "id": f"chapter_{roman_num}",
#             "name": capitalize_title(full_chapter_title),  # Gán tiêu đề đầy đủ với số La Mã chuẩn
#             "number": roman_num,
#             "articles": []
#         }
#         law["chapters"].append(current_chapter)
#         current_section = None
#         current_article = None
#         current_clause = None
#         current_point = None
#         i = j  # Cập nhật chỉ số sau khi xử lý tiêu đề chương
#         continue

#     # Kiểm tra tiêu đề Mục
#     section_match = section_pattern.match(line)
#     if section_match and current_chapter:
#         section_number = section_match.group(1)
#         section_title = section_match.group(2).strip()
#         full_section_header = f"Mục {section_number} {section_title}".strip()
#         j = i + 1
#         while j < len(lines):
#             next_line = lines[j].strip()
#             if not next_line or chapter_pattern.match(next_line) or article_pattern.match(next_line) or section_pattern.match(next_line):
#                 break
#             full_section_header += " " + next_line
#             j += 1
#         current_section = capitalize_title(full_section_header)
#         i = j
#         continue

#     # Kiểm tra tiêu đề Điều
#     article_match = article_pattern.match(line)
#     if article_match and current_chapter:
#         article_number = article_match.group(1)
#         article_title = article_match.group(2).strip()
#         current_article = {
#             "id": f"article_{article_number}",
#             "parent": current_section if current_section else None,
#             "name": capitalize_title(article_title),
#             "number": int(article_number),
#             "content": "",
#             "clauses": [],
#             "references": []
#         }
#         current_chapter["articles"].append(current_article)
#         current_clause = None
#         current_point = None
#         i += 1
#         continue

#     # Kiểm tra tiêu đề của điều khoản (clause)
#     clause_match = clause_pattern.match(line)
#     if clause_match and current_article:
#         clause_number = clause_match.group(1)
#         clause_text = clause_match.group(2).strip()
#         current_clause = {
#             "id": f"clause_{current_article['number']}_{clause_number}",
#             "number": int(clause_number),
#             "content": clause_text,
#             "points": []
#         }
#         current_article["clauses"].append(current_clause)
#         current_point = None
#         i += 1
#         continue

#     # Kiểm tra các điểm (points) trong khoản
#     point_match = point_pattern.match(line)
#     if point_match and current_clause:
#         point_label = point_match.group(1)
#         point_text = point_match.group(2).strip()
#         current_point = {
#             "id": f"point_{current_article['number']}_{current_clause['number']}_{point_label}",
#             "number": point_label,
#             "content": point_text
#         }
#         current_clause["points"].append(current_point)
#         i += 1
#         continue

#     # Xử lý tham chiếu chéo
#     reference_match = re.search(r"Điều\s+(\d+)", line)
#     if reference_match and current_article and not clause_match and not point_match:
#         ref_article = reference_match.group(1)
#         if int(ref_article) != current_article["number"] and ref_article not in current_article["references"]:
#             current_article["references"].append(f"article_{ref_article}")

#     # Gán nội dung vào cấp phù hợp
#     if current_point:
#         current_point["content"] += " " + line
#     elif current_clause:
#         current_clause["content"] += " " + line
#     elif current_article:
#         current_article["content"] += " " + line

#     i += 1

# # Lưu kết quả ra file JSON
# with open("law_structure.json", "w", encoding="utf-8") as f:
#     json.dump({"law": law}, f, ensure_ascii=False, indent=4)

# print("Đã chuyển file DOCX sang cấu trúc JSON với tiêu đề chương đầy đủ và số La Mã chuẩn hóa thành công.")







########################################


import re
import json
from docx import Document

# Đọc dữ liệu từ file DOCX
doc = Document("31-2024-qh15_1.docx")
lines = [para.text for para in doc.paragraphs]

# Khởi tạo cấu trúc cơ bản của văn bản luật
law = {
    "id": "law_001",
    "title": "LUẬT ĐẤT ĐAI",
    "date": "01-03-2024",
    "source": "Quốc hội",
    "category": "Đất đai",
    "chapters": []
}

# Định nghĩa các biểu thức chính quy
chapter_pattern = re.compile(r"^Chương\s+([IVXLCDM]+)(.*)", re.IGNORECASE)
section_pattern = re.compile(r"^Mục\s+(\d+)\s*(.*)", re.IGNORECASE)
article_pattern = re.compile(r"^Điều\s+(\d+)[\.\:]\s*(.*)", re.IGNORECASE)
clause_pattern = re.compile(r"^(\d+)\.\s*(.*)")
point_pattern = re.compile(r"^([a-z])[)]\s*(.*)")

# Hàm chuẩn hóa tiêu đề với xử lý số La Mã
def capitalize_title(text):
    words = text.split()
    for i, word in enumerate(words):
        if re.match(r'^[IVXLCDM]+$', word, re.IGNORECASE):
            words[i] = word.upper()
        else:
            words[i] = word.capitalize()
    return " ".join(words)

i = 0
current_chapter = None
current_section = None
current_article = None
current_clause = None
current_point = None

while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue

    # Kiểm tra tiêu đề Chương
    chapter_match = chapter_pattern.match(line)
    if chapter_match:
        roman_num = chapter_match.group(1).upper()
        chapter_title = chapter_match.group(2).strip()
        full_chapter_title = f"Chương {roman_num} {chapter_title}".strip()
        
        j = i + 1
        while j < len(lines):
            next_line = lines[j].strip()
            if not next_line or chapter_pattern.match(next_line) or section_pattern.match(next_line) or article_pattern.match(next_line):
                break
            full_chapter_title += " " + next_line
            j += 1
        
        current_chapter = {
            "id": f"chapter_{roman_num}",
            "name": capitalize_title(full_chapter_title),
            "number": roman_num,
            "articles": []
        }
        law["chapters"].append(current_chapter)
        current_section = None
        current_article = None
        current_clause = None
        current_point = None
        i = j
        continue

    # Kiểm tra tiêu đề Mục
    section_match = section_pattern.match(line)
    if section_match and current_chapter:
        section_number = section_match.group(1)
        section_title = section_match.group(2).strip()
        full_section_header = f"Mục {section_number} {section_title}".strip()
        j = i + 1
        while j < len(lines):
            next_line = lines[j].strip()
            if not next_line or chapter_pattern.match(next_line) or section_pattern.match(next_line) or article_pattern.match(next_line):
                break
            full_section_header += " " + next_line
            j += 1
        current_section = capitalize_title(full_section_header)
        i = j
        continue

    # Kiểm tra tiêu đề Điều
    article_match = article_pattern.match(line)
    if article_match and current_chapter:
        article_number = article_match.group(1)
        article_title = article_match.group(2).strip()
        current_article = {
            "id": f"article_{article_number}",
            "parent": current_section if current_section else None,
            "name": capitalize_title(article_title),
            "number": int(article_number),
            "content": "",
            "clauses": [],
            "references": []
        }
        current_chapter["articles"].append(current_article)
        current_clause = None
        current_point = None
        i += 1
        
        # Thu thập nội dung tổng quan của điều luật
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line or chapter_pattern.match(next_line) or section_pattern.match(next_line) or article_pattern.match(next_line) or clause_pattern.match(next_line):
                break
            current_article["content"] += " " + next_line
            i += 1
        current_article["content"] = current_article["content"].strip()
        continue

    # Kiểm tra tiêu đề của điều khoản (clause)
    clause_match = clause_pattern.match(line)
    if clause_match and current_article:
        clause_number = clause_match.group(1)
        clause_text = clause_match.group(2).strip()
        current_clause = {
            "id": f"clause_{current_article['number']}_{clause_number}",
            "number": int(clause_number),
            "content": clause_text,
            "points": []
        }
        current_article["clauses"].append(current_clause)
        current_point = None
        i += 1
        continue

    # Kiểm tra các điểm (points) trong khoản
    point_match = point_pattern.match(line)
    if point_match and current_clause:
        point_label = point_match.group(1)
        point_text = point_match.group(2).strip()
        current_point = {
            "id": f"point_{current_article['number']}_{current_clause['number']}_{point_label}",
            "number": point_label,
            "content": point_text
        }
        current_clause["points"].append(current_point)
        i += 1
        continue

    # Xử lý tham chiếu chéo
    reference_match = re.search(r"(?:Điều\s+(\d+)|khoản\s+(\d+)\s+Điều\s+(\d+))", line)
    if reference_match and current_article and not clause_match and not point_match:
        if reference_match.group(1):  # "Điều X"
            ref_article = reference_match.group(1)
            if int(ref_article) != current_article["number"] and f"article_{ref_article}" not in current_article["references"]:
                current_article["references"].append(f"article_{ref_article}")
        elif reference_match.group(2) and reference_match.group(3):  # "khoản X Điều Y"
            ref_article = reference_match.group(3)
            if int(ref_article) != current_article["number"] and f"article_{ref_article}" not in current_article["references"]:
                current_article["references"].append(f"article_{ref_article}")

    # Gán nội dung vào cấp phù hợp
    if current_point:
        current_point["content"] += " " + line
    elif current_clause:
        current_clause["content"] += " " + line
    elif current_article:
        current_article["content"] += " " + line

    i += 1

# Lưu kết quả ra file JSON
with open("law_structure.json", "w", encoding="utf-8") as f:
    json.dump({"law": law}, f, ensure_ascii=False, indent=4)

print("Đã cập nhật cấu trúc JSON với sections và cải thiện xử lý content và references.")